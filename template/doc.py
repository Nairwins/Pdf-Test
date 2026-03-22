import json
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


# ── Units ─────────────────────────────────────────────────────────────────────
def twip(pt):  return int(round(pt * 20))
def emu(pt):   return int(round(pt * 12700))
def hx(c):     return c.lstrip("#").upper()


# ── XML run ───────────────────────────────────────────────────────────────────
def _run(para, text, font, size_pt, bold=False, italic=False, color="#111111"):
    r   = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf  = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), font)
    rPr.append(rf)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    if italic:
        rPr.append(OxmlElement("w:i"))
        rPr.append(OxmlElement("w:iCs"))
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), str(int(size_pt * 2)))
        rPr.append(el)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), hx(color))
    rPr.append(col)
    # Disable kerning to match ReportLab's plain rendering
    kern = OxmlElement("w:kern")
    kern.set(qn("w:val"), "0")
    rPr.append(kern)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    para._p.append(r)


def _tab_run(para):
    r = OxmlElement("w:r")
    r.append(OxmlElement("w:tab"))
    para._p.append(r)


# ── Paragraph helpers ─────────────────────────────────────────────────────────
def _spacing(para, before=0, after=0, line_pt=None, rule="exact"):
    pPr = para._p.get_or_add_pPr()
    sp  = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:before"), str(twip(before)))
    sp.set(qn("w:after"),  str(twip(after)))
    if line_pt is not None:
        sp.set(qn("w:line"),     str(twip(line_pt)))
        sp.set(qn("w:lineRule"), rule)


def _indent(para, left_pt=0, hanging_pt=0):
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if left_pt:
        ind.set(qn("w:left"),    str(twip(left_pt)))
    if hanging_pt:
        ind.set(qn("w:hanging"), str(twip(hanging_pt)))


def _border_bottom(para, color, size_pt, space_pt=0):
    pPr  = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    str(int(size_pt * 8)))
    b.set(qn("w:space"), str(int(space_pt)))
    b.set(qn("w:color"), hx(color))
    pBdr.append(b)


# CW = 595.28 - 52 - 52 = 491.28 pt
_CW_TWIP = twip(491.28)


def _tab_right(para):
    pPr  = para._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    t = OxmlElement("w:tab")
    t.set(qn("w:val"), "right")
    t.set(qn("w:pos"), str(_CW_TWIP))
    tabs.append(t)


def _no_ctx(para):
    pPr = para._p.get_or_add_pPr()
    cs  = OxmlElement("w:contextualSpacing")
    cs.set(qn("w:val"), "0")
    pPr.append(cs)


def _bullet_numpr(para):
    pPr   = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl  = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); numPr.append(ilvl)
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "1"); numPr.append(numId)
    first = pPr.find(qn("w:spacing")) or pPr.find(qn("w:ind"))
    if first is not None:
        pPr.insert(list(pPr).index(first), numPr)
    else:
        pPr.insert(0, numPr)


def _zero_normal_style(doc):
    for pStyle in doc.styles.element.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style"):
        sid = pStyle.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId", "")
        if sid == "Normal":
            pPr = pStyle.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
            if pPr is None:
                pPr = OxmlElement("w:pPr"); pStyle.append(pPr)
            sp = pPr.find(qn("w:spacing"))
            if sp is None:
                sp = OxmlElement("w:spacing"); pPr.append(sp)
            sp.set(qn("w:before"),   "0")
            sp.set(qn("w:after"),    "0")
            sp.set(qn("w:line"),     str(twip(10)))
            sp.set(qn("w:lineRule"), "exact")
            break


# ─────────────────────────────────────────────────────────────────────────────
def build_resume_docx(data, main_color="#1a56db", secondary_color="#6b7280", font="Helvetica"):
    if isinstance(data, (str, Path)):
        with open(data, "r", encoding="utf-8") as f:
            data = json.load(f)

    FONT    = "Arial" if font == "Helvetica" else font
    C_ACC   = main_color
    C_META  = secondary_color
    C_BLACK = "#111111"
    C_BODY  = "#222222"

    doc = Document()
    _zero_normal_style(doc)

    for sec in doc.sections:
        sec.page_width    = emu(595.28)
        sec.page_height   = emu(841.89)
        sec.left_margin   = emu(52)
        sec.right_margin  = emu(52)
        sec.top_margin    = emu(50)
        sec.bottom_margin = emu(36)

    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # ── Para factory ─────────────────────────────────────────────────────────
    def P():
        p = doc.add_paragraph()
        _no_ctx(p)
        _spacing(p, 0, 0)
        return p

    # ── Spacer ───────────────────────────────────────────────────────────────
    # Implemented as after-spacing on the PREVIOUS para to avoid extra para overhead.
    # But since we don't always know the previous, we use a minimal empty para.
    def spacer(h):
        p = P()
        _spacing(p, 0, 0, line_pt=h, rule="exact")

    # ── Section header ────────────────────────────────────────────────────────
    # PDF: draw 8pt bold text at y → y-=3 → draw rule → y-=10
    # Baseline-to-baseline from section header to first content = 3+10 = 13pt
    # DOCX: single para, line=11pt (close to 8pt font + small border gap),
    #        after=2pt (extra space after border to reach ~13pt total)
    # Calibrated: line=11 + after=2 → total ~13pt from hdr top to content top.
    # The border is on the bottom of this para (space_pt=3 = 3pt gap text→border).
    def section(label):
        p = P()
        _run(p, label.upper(), FONT, 8, bold=True, color=C_ACC)
        _spacing(p, before=0, after=0, line_pt=6.9, rule="exact")
        _border_bottom(p, C_ACC, size_pt=0.8, space_pt=5)

    # ── Two-column (left + right-aligned date) ────────────────────────────────
    # PDF: drawString left + drawString right (at ML+CW-rw)
    # DOCX: right tab stop at CW
    def two_col(left, right, left_size=10, right_size=8,
                left_bold=True, left_color=C_BLACK, right_color=C_META, line_pt=13):
        p = P()
        _tab_right(p)
        _run(p, left,  FONT, left_size, bold=left_bold, color=left_color)
        _tab_run(p)
        _run(p, right, FONT, right_size, color=right_color)
        _spacing(p, 0, 0, line_pt=line_pt, rule="exact")

    # ── Subline (italic meta) ─────────────────────────────────────────────────
    # PDF: FI 8.5pt, y -= 11
    def subline(text):
        p = P()
        _run(p, text, FONT, 8.5, italic=True, color=C_META)
        _spacing(p, 0, 0, line_pt=11, rule="exact")

    # ── Bullet item ───────────────────────────────────────────────────────────
    # PDF: circle bullet at ML+7, text at ML+12, lh=11.5
    def bullet(text):
        p = P()
        _bullet_numpr(p)
        _run(p, text, FONT, 8.5, color=C_BODY)
        _spacing(p, 0, 0, line_pt=11.5, rule="exact")
        _indent(p, left_pt=12, hanging_pt=5)

    # ═══════════════════════════════════════════════════════════════════════════
    # HEADER
    # ═══════════════════════════════════════════════════════════════════════════
    name  = data.get("name",  "")
    title = data.get("title", "")
    c_    = data.get("contact", {})

    # Name: 26pt bold. PDF y-=16.
    # Calibration: line=16 gives title 2.9pt too high → after=3pt compensates.
    p = P()
    _run(p, name, FONT, 26, bold=True, color=C_BLACK)
    _spacing(p, before=0, after=3, line_pt=16, rule="exact")

    # Title: 10pt accent. PDF y-=13.
    # Calibration: line=12 (was 13, reduced by 1 to fix 1.1pt over-gap to contact).
    if title:
        p = P()
        _run(p, title, FONT, 10, color=C_ACC)
        _spacing(p, 0, 0, line_pt=12, rule="exact")

    # Contact row: 7.8pt meta. PDF y-=14.
    # Calibration: line=14 gives rule 6.9pt too high → after=7pt compensates.
    ck = [("email","E"),("phone","P"),("location","L"),
          ("github","G"),("website","W"),("linkedin","Li")]
    parts = [c_.get(k, "") for k, _ in ck if c_.get(k, "")]
    if parts:
        p = P()
        for i, val in enumerate(parts):
            _run(p, val, FONT, 7.8, color=C_META)
            if i < len(parts) - 1:
                _run(p, "   \u00b7   ", FONT, 7.8, color=C_META)
        _spacing(p, before=0, after=7, line_pt=14, rule="exact")

    # Accent rule under header: PDF cv.line then y-=14.
    # Calibration: border para line=2pt (nearly invisible), after=12pt.
    # Total: 2+12=14pt gap, matching PDF's y-=14.
    p = P()
    _border_bottom(p, C_ACC, size_pt=1.8, space_pt=1)
    _spacing(p, before=0, after=5, line_pt=2, rule="exact")

    # ═══════════════════════════════════════════════════════════════════════════
    # SUMMARY
    # ═══════════════════════════════════════════════════════════════════════════
    summary = data.get("summary", "").strip()
    if summary:
        section("Summary")
        p = P()
        _run(p, summary, FONT, 9, italic=True, color=C_BODY)
        _spacing(p, 0, 0, line_pt=12, rule="exact")
        spacer(4)

    # ═══════════════════════════════════════════════════════════════════════════
    # EXPERIENCE
    # ═══════════════════════════════════════════════════════════════════════════
    experience = data.get("experience", [])
    if experience:
        section("Experience")
        for exp in experience:
            role     = exp.get("role",        "")
            company  = exp.get("company",     "")
            loc      = exp.get("location",    "")
            start    = exp.get("start_date",  "")
            end      = exp.get("end_date",    "Present")
            descs    = exp.get("description", [])
            date_str = f"{start} - {end}" if start else end
            comp_str = company + (f"  \u00b7  {loc}" if loc else "")
            two_col(role, date_str)
            subline(comp_str)
            for d in descs:
                bullet(d)
            spacer(7)

    # ═══════════════════════════════════════════════════════════════════════════
    # PROJECTS
    # ═══════════════════════════════════════════════════════════════════════════
    projects = data.get("projects", [])
    if projects:
        section("Projects")
        for proj in projects:
            pname = proj.get("name",  "")
            links = proj.get("links", "")
            descs = proj.get("description", [])
            if links:
                two_col(pname, links, left_size=9.5, right_size=8,
                        left_bold=True, right_color=C_META)
            else:
                p = P()
                _run(p, pname, FONT, 9.5, bold=True, color=C_BLACK)
                _spacing(p, 0, 0, line_pt=13, rule="exact")
            for d in descs:
                bullet(d)
            spacer(6)

    # ═══════════════════════════════════════════════════════════════════════════
    # SKILLS
    # ═══════════════════════════════════════════════════════════════════════════
    hard = data.get("hardskills", {})
    soft = data.get("softskills", [])
    if hard or soft:
        section("Skills")
        items = hard if isinstance(hard, dict) else ({"Technical": hard} if hard else {})
        for cat, vals in items.items():
            val_str = ", ".join(vals) if isinstance(vals, list) else str(vals)
            p = P()
            _run(p, cat + ": ", FONT, 8.5, bold=True, color=C_BLACK)
            _run(p, val_str,    FONT, 8.5, color=C_BODY)
            _spacing(p, 0, 0, line_pt=12, rule="exact")
        if soft:
            p = P()
            _run(p, "Soft: ",        FONT, 8.5, bold=True, color=C_BLACK)
            _run(p, ", ".join(soft), FONT, 8.5, color=C_BODY)
            _spacing(p, 0, 0, line_pt=12, rule="exact")
        spacer(4)

    # ═══════════════════════════════════════════════════════════════════════════
    # EDUCATION
    # ═══════════════════════════════════════════════════════════════════════════
    education = data.get("education", [])
    if education:
        section("Education")
        for edu in education:
            deg      = edu.get("degree",      "")
            field    = edu.get("field",       "")
            inst     = edu.get("institution", "")
            start    = edu.get("start_date",  "")
            end      = edu.get("end_date",    "")
            state    = edu.get("state",       "")
            gpa      = edu.get("gpa",         "")
            deg_str  = deg + (f" - {field}" if field else "")
            date_str = (f"{start} - " if start else "") + (end or "")
            if state:
                date_str += f"  [{state}]"
            meta = "  \u00b7  ".join(
                x for x in [inst, f"GPA: {gpa}" if gpa else ""] if x)
            two_col(deg_str, date_str)
            subline(meta)
            spacer(5)

    # ═══════════════════════════════════════════════════════════════════════════
    # CERTIFICATIONS
    # ═══════════════════════════════════════════════════════════════════════════
    certs = data.get("certifications", [])
    if certs:
        section("Certifications")
        for cert in certs:
            bullet(str(cert))
        spacer(4)

    # ═══════════════════════════════════════════════════════════════════════════
    # LANGUAGES
    # ═══════════════════════════════════════════════════════════════════════════
    langs = data.get("languages", [])
    if langs:
        section("Languages")
        p = P()
        _run(p, "  \u00b7  ".join(langs), FONT, 8.5, color=C_BODY)
        _spacing(p, 0, 0, line_pt=12, rule="exact")
        spacer(3)

    # ═══════════════════════════════════════════════════════════════════════════
    # AWARDS / PUBLICATIONS
    # ═══════════════════════════════════════════════════════════════════════════
    for key, label in [("awards", "Awards"), ("publications", "Publications")]:
        items = data.get(key, [])
        if items:
            section(label)
            for item in items:
                bullet(str(item))
            spacer(4)

    _inject_numbering(doc, C_ACC)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _inject_numbering(doc, accent_color):
    xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="0">
    <w:multiLevelType w:val="hybridMultilevel"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="&#x25CF;"/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:left="{twip(12)}" w:hanging="{twip(5)}"/></w:pPr>
      <w:rPr>
        <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
        <w:color w:val="{hx(accent_color)}"/>
        <w:sz w:val="10"/><w:szCs w:val="10"/>
      </w:rPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
</w:numbering>""".encode("utf-8")

    np = doc.part.numbering_part
    if np is not None:
        np._element = etree.fromstring(xml)
    else:
        from docx.parts.numbering import NumberingPart
        np = NumberingPart.new()
        np._element = etree.fromstring(xml)
        doc.part.relate_to(
            np,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering")
